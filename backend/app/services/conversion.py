"""Conversão PDF/DOCX/TXT -> texto UTF-8 (REQ-FUNC-002)."""
import logging
from dataclasses import dataclass, field
from pathlib import Path

from app.services.ingestion import sha256_text

logger = logging.getLogger(__name__)


class ConversionError(Exception):
    pass


@dataclass
class ConvertedDocument:
    text: str
    page_count: int
    word_count: int
    sha256: str
    # Lista de (número_da_página, texto_da_página). None para formatos sem paginação.
    pages: list[tuple[int | None, str]] = field(default_factory=list)


def _finalize(pages: list[tuple[int | None, str]]) -> ConvertedDocument:
    text = "\n\n".join(t.strip() for _, t in pages if t.strip())
    numbered = [p for p in pages if p[0] is not None]
    page_count = (max(p for p, _ in numbered) if numbered else 1) if text else 0
    return ConvertedDocument(
        text=text,
        page_count=page_count,
        word_count=len(text.split()),
        sha256=sha256_text(text),
        pages=pages,
    )


def convert_pdf(path: Path) -> ConvertedDocument:
    """PDF -> texto preservando a ordem dos parágrafos; rastreia páginas por trecho."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        pages: list[tuple[int | None, str]] = []
        for i, page in enumerate(reader.pages, start=1):
            pages.append((i, page.extract_text() or ""))
    except Exception as exc:  # noqa: BLE001 — PDF corrompido vira erro de conversão
        raise ConversionError(f"Falha ao converter PDF: {exc}") from exc
    doc = _finalize(pages)
    doc.page_count = len(reader.pages)
    return doc


def convert_docx(path: Path) -> ConvertedDocument:
    """DOCX -> apenas parágrafos do corpo (sem cabeçalho/rodapé/notas)."""
    try:
        import docx

        document = docx.Document(str(path))
        body_text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    except ConversionError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ConversionError(f"Falha ao converter DOCX: {exc}") from exc
    return _finalize([(None, body_text)])


def convert_txt(path: Path) -> ConvertedDocument:
    """TXT -> leitura direta em UTF-8."""
    try:
        text = path.read_bytes().decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        raise ConversionError(f"Falha ao ler TXT: {exc}") from exc
    return _finalize([(None, text)])


def convert(path: Path) -> ConvertedDocument:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return convert_pdf(path)
    if ext == ".docx":
        return convert_docx(path)
    if ext == ".txt":
        return convert_txt(path)
    raise ConversionError(f"Extensão não suportada para conversão: {ext}")
